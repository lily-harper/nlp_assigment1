#!/usr/bin/env python3
"""
Count words in a corpus of .docx and .pdf files using basic whitespace
tokenization.

Install dependencies:
    pip install python-docx pypdf

Examples:
    python count_words.py ./documents
    python count_words.py ./documents --top 25
    python count_words.py file1.pdf file2.docx --output counts.csv
"""

from __future__ import annotations

import argparse
import csv
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


def extract_docx_text(file_path: Path) -> str:
    """Extract text from paragraphs and tables in a Word document."""
    document = Document(file_path)
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def extract_pdf_text(file_path: Path) -> str:
    """Extract text from a text-based PDF."""
    reader = PdfReader(file_path)
    text_parts: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_text(file_path: Path) -> str:
    """Choose the appropriate extraction method based on file extension."""
    extension = file_path.suffix.lower()

    if extension == ".docx":
        return extract_docx_text(file_path)

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    raise ValueError(f"Unsupported file type: {extension}")


def tokenize(text: str, lowercase: bool = False) -> list[str]:
    """
    Tokenize text using whitespace only.

    Because this uses str.split(), punctuation remains attached:
    'word', 'word,' and 'word.' are counted as different tokens.
    """
    if lowercase:
        text = text.lower()

    return text.split()


def collect_files(input_paths: list[Path], recursive: bool) -> list[Path]:
    """Collect supported files from the supplied files and directories."""
    collected: set[Path] = set()

    for input_path in input_paths:
        if input_path.is_file():
            if input_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                collected.add(input_path.resolve())
            else:
                print(
                    f"Skipping unsupported file: {input_path}",
                    file=sys.stderr,
                )

        elif input_path.is_dir():
            pattern = "**/*" if recursive else "*"

            for candidate in input_path.glob(pattern):
                if (
                    candidate.is_file()
                    and candidate.suffix.lower() in SUPPORTED_EXTENSIONS
                    and not candidate.name.startswith("~$")
                ):
                    collected.add(candidate.resolve())

        else:
            print(f"Path not found: {input_path}", file=sys.stderr)

    return sorted(collected)


def write_counts_to_csv(word_counts: Counter[str], output_path: Path) -> None:
    """Write corpus-wide word frequencies to a CSV file."""
    with output_path.open("w", newline="", encoding="utf-8") as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["word", "count"])

        for word, count in word_counts.most_common():
            writer.writerow([word, count])


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Count words in .docx and .pdf files using whitespace tokenization."
        )
    )

    parser.add_argument(
        "paths",
        nargs="+",
        type=Path,
        help="One or more files or directories containing corpus documents.",
    )

    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Search directories and their subdirectories.",
    )

    parser.add_argument(
        "--lowercase",
        action="store_true",
        help="Convert text to lowercase before counting.",
    )

    parser.add_argument(
        "--top",
        type=int,
        default=20,
        help="Number of most common words to display. Default: 20.",
    )

    parser.add_argument(
        "--output",
        type=Path,
        help="Optional CSV file for all corpus-wide word counts.",
    )

    return parser.parse_args()


def main() -> int:
    args = parse_arguments()
    files = collect_files(args.paths, args.recursive)

    if not files:
        print("No supported .docx or .pdf files were found.", file=sys.stderr)
        return 1

    corpus_counts: Counter[str] = Counter()
    total_words = 0
    successful_files = 0

    print("\nPer-file word counts")
    print("-" * 60)

    for file_path in files:
        try:
            text = extract_text(file_path)
            tokens = tokenize(text, lowercase=args.lowercase)
        except Exception as error:
            print(
                f"Could not process {file_path}: {error}",
                file=sys.stderr,
            )
            continue

        word_count = len(tokens)
        total_words += word_count
        corpus_counts.update(tokens)
        successful_files += 1

        print(f"{word_count:>10,}  {file_path}")

    if successful_files == 0:
        print("None of the files could be processed.", file=sys.stderr)
        return 1

    print("\nCorpus summary")
    print("-" * 60)
    print(f"Files processed: {successful_files:,}")
    print(f"Total words:     {total_words:,}")
    print(f"Unique tokens:   {len(corpus_counts):,}")

    if args.top > 0:
        print(f"\nTop {args.top} tokens")
        print("-" * 60)

        for word, count in corpus_counts.most_common(args.top):
            print(f"{count:>10,}  {word}")

    if args.output:
        try:
            write_counts_to_csv(corpus_counts, args.output)
            print(f"\nFull frequency table saved to: {args.output}")
        except OSError as error:
            print(f"Could not write CSV: {error}", file=sys.stderr)
            return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())