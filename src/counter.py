#!/usr/bin/env python3
"""
Count words in a corpus of .docx and .pdf files using basic whitespace
tokenization.

Install dependencies:
    pip install python-docx pypdf reportlab

Examples:
    python counter.py
    python counter.py --paths-file another_paths_file.txt
    python counter.py file1.pdf file2.docx --output counts.csv
"""

from __future__ import annotations

import argparse
import csv
import sys
import textwrap
from collections import Counter
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
DEFAULT_PATHS_FILE = Path(__file__).resolve().parent / "corpus" / "corpus_paths.txt"


def read_paths_file(paths_file: Path) -> list[Path]:
    """Read file and directory paths from a text file."""
    paths: list[Path] = []

    with paths_file.open("r", encoding="utf-8") as file:
        for line_number, line in enumerate(file, start=1):
            line = line.strip()

            # Ignore blank lines and comments.
            if not line or line.startswith("#"):
                continue

            path = Path(line).expanduser()

            # Resolve relative paths from the directory containing the paths file.
            if not path.is_absolute():
                path = paths_file.parent / path

            if not path.exists():
                print(
                    f"Warning: path on line {line_number} does not exist: {path}",
                    file=sys.stderr,
                )
                continue

            paths.append(path)

    return paths

def extract_docx_text(file_path: Path) -> str:
    """Extract text from paragraphs and tables in a Word document."""
    document = Document(file_path)
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def extract_pdf_text(file_path: Path) -> str:
    """Extract text from a text-based PDF."""
    reader = PdfReader(file_path)
    text_parts: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_text(file_path: Path) -> str:
    """Choose the appropriate extraction method based on file extension."""
    extension = file_path.suffix.lower()

    if extension == ".docx":
        return extract_docx_text(file_path)

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    raise ValueError(f"Unsupported file type: {extension}")


def tokenize(text: str, lowercase: bool = False) -> list[str]:
    """
    Tokenize text using whitespace only.

    Because this uses str.split(), punctuation remains attached:
    'word', 'word,' and 'word.' are counted as different tokens.
    """
    if lowercase:
        text = text.lower()

    return text.split()


def collect_files(input_paths: list[Path]) -> list[Path]:
    """Collect supported files from the supplied files and directories."""
    collected: set[Path] = set()

    for input_path in input_paths:
        if input_path.is_file():
            if input_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                collected.add(input_path.resolve())
            else:
                print(
                    f"Skipping unsupported file: {input_path}",
                    file=sys.stderr,
                )

        elif input_path.is_dir():
            for candidate in input_path.rglob("*"):
                if (
                    candidate.is_file()
                    and candidate.suffix.lower() in SUPPORTED_EXTENSIONS
                    and not candidate.name.startswith("~$")
                ):
                    collected.add(candidate.resolve())

        else:
            print(f"Path not found: {input_path}", file=sys.stderr)

    return sorted(collected)


def write_counts_to_csv(word_counts: Counter[str], output_path: Path) -> None:
    """Write corpus-wide word frequencies to a CSV file."""
    with output_path.open("w", newline="", encoding="utf-8") as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["word", "count"])

        for word, count in word_counts.most_common():
            writer.writerow([word, count])


def build_report(
    file_results: list[tuple[Path, int, int]],
    total_words: int,
    corpus_counts: Counter[str],
) -> list[str]:
    """Build the lines used by both text and PDF reports."""
    lines = [
        "Corpus Word Count Report",
        "=" * 60,
        f"Files processed: {len(file_results):,}",
        f"Total words: {total_words:,}",
        f"Total unique words: {len(corpus_counts):,}",
        "",
        "Files",
        "-" * 60,
    ]

    for file_path, word_count, unique_count in file_results:
        lines.append(str(file_path))
        lines.append(
            f"  Words: {word_count:,} | Unique words: {unique_count:,}"
        )

    lines.extend(["", "Top 20 tokens", "-" * 60])
    lines.extend(
        f"{count:>10,}  {word}"
        for word, count in corpus_counts.most_common(20)
    )
    return lines


def write_report(lines: list[str], output_path: Path) -> None:
    """Write a corpus report as a text file or PDF."""
    extension = output_path.suffix.lower()

    if extension == ".txt":
        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return

    if extension == ".pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError as error:
            raise RuntimeError(
                "PDF reports require reportlab: pip install reportlab"
            ) from error

        pdf = canvas.Canvas(str(output_path), pagesize=letter)
        _, page_height = letter
        left_margin = 50
        top_margin = 50
        line_height = 14
        y_position = page_height - top_margin

        for line in lines:
            wrapped_lines = textwrap.wrap(line, width=90) or [""]

            for wrapped_line in wrapped_lines:
                if y_position < top_margin:
                    pdf.showPage()
                    y_position = page_height - top_margin

                printable_line = wrapped_line.encode(
                    "latin-1", errors="replace"
                ).decode("latin-1")
                pdf.drawString(left_margin, y_position, printable_line)
                y_position -= line_height

        pdf.save()
        return

    raise ValueError("Report filename must end in .txt or .pdf")


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Count words in .docx and .pdf files using whitespace tokenization."
        )
    )

    parser.add_argument(
        "--paths-file",
        type=Path,
        default=DEFAULT_PATHS_FILE,
        help=(
            "Text file containing one corpus file or directory path per line. "
            "Default: corpus/corpus_paths.txt."
        ),
    )

    parser.add_argument(
        "paths",
        nargs="*",
        type=Path,
        help="Optional additional corpus files or directories.",
    )

    parser.add_argument(
        "--lowercase",
        action="store_true",
        help="Convert text to lowercase before counting.",
    )

    parser.add_argument(
        "--output",
        type=Path,
        help="Optional CSV file for all corpus-wide word counts.",
    )

    parser.add_argument(
        "--report",
        type=Path,
        help="Save a summary report as a .txt or .pdf file.",
    )

    return parser.parse_args()


def main() -> int:
    args = parse_arguments()
    input_paths = list(args.paths)

    if args.paths_file:
        try:
            input_paths.extend(read_paths_file(args.paths_file))
        except OSError as error:
            print(f"Could not read paths file {args.paths_file}: {error}", file=sys.stderr)
            return 1

    files = collect_files(input_paths)

    if not files:
        print("No supported .docx or .pdf files were found.", file=sys.stderr)
        return 1

    corpus_counts: Counter[str] = Counter()
    total_words = 0
    successful_files = 0
    file_results: list[tuple[Path, int, int]] = []

    for file_path in files:
        print(f"\nCounting words in {file_path}")

        try:
            text = extract_text(file_path)
            tokens = tokenize(text, lowercase=args.lowercase)
        except Exception as error:
            print(
                f"Could not process {file_path}: {error}",
                file=sys.stderr,
            )
            continue

        word_count = len(tokens)
        total_words += word_count
        corpus_counts.update(tokens)
        successful_files += 1
        unique_count = len(set(tokens))
        file_results.append((file_path, word_count, unique_count))

        print(f"Word count: {word_count:,}")
        print(f"Unique words in document: {unique_count:,}")

    if successful_files == 0:
        print("None of the files could be processed.", file=sys.stderr)
        return 1

    print("\nCorpus summary")
    print("-" * 60)
    print(f"Files processed: {successful_files:,}")
    print(f"Total words:     {total_words:,}")
    print(f"Total unique words: {len(corpus_counts):,}")

    print("\nTop 20 tokens")
    print("-" * 60)

    for word, count in corpus_counts.most_common(20):
        print(f"{count:>10,}  {word}")

    if args.report:
        try:
            report_lines = build_report(file_results, total_words, corpus_counts)
            write_report(report_lines, args.report)
            print(f"\nReport saved to: {args.report}")
        except (OSError, RuntimeError, ValueError) as error:
            print(f"Could not write report: {error}", file=sys.stderr)
            return 1

    if args.output:
        try:
            write_counts_to_csv(corpus_counts, args.output)
            print(f"\nFull frequency table saved to: {args.output}")
        except OSError as error:
            print(f"Could not write CSV: {error}", file=sys.stderr)
            return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
